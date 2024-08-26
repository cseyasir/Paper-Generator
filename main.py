import streamlit as st
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
import io
from docx import Document


def inject_custom_css():
    st.markdown("""
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Roboto:wght@400;700&display=swap');

    /* Apply global styles */
    body {
        font-family: 'Roboto', sans-serif;
        background-color: #f7f9fc;
        color: #333333;
    }

    /* Header Styles */
    h1, h2, h3, h4, h5, h6 {
        font-weight: 700;
        color: #2c3e50;
    }

    /* Main Container Padding */
    .main {
        padding: 20px;
    }

    /* Button Styles */
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        padding: 12px 24px;
        border: none;
        border-radius: 8px;
        font-size: 16px;
        font-weight: 600;
        transition: background-color 0.3s, transform 0.2s;
    }
    .stButton > button:hover {
        background-color: #45a049;
        transform: scale(1.05);
    }

    /* Download Button Styles */
    .stDownloadButton > button {
        background-color: #008CBA;
        color: white;
        padding: 12px 24px;
        border: none;
        border-radius: 8px;
        font-size: 16px;
        font-weight: 600;
        transition: background-color 0.3s, transform 0.2s;
    }
    .stDownloadButton > button:hover {
        background-color: #007bb5;
        transform: scale(1.05);
    }

    /* Text Input Styles */
    .stTextInput > div > div > input {
        border: 2px solid #ccc;
        border-radius: 8px;
        padding: 10px;
        font-size: 14px;
        transition: border-color 0.3s;
    }
    .stTextInput > div > div > input:focus {
        border-color: #4CAF50;
        box-shadow: 0 0 5px rgba(76, 175, 80, 0.5);
    }

    /* Text Area Styles */
    .stTextArea > div > div > textarea {
        border: 2px solid #ccc;
        border-radius: 8px;
        padding: 10px;
        font-size: 14px;
        transition: border-color 0.3s;
    }
    .stTextArea > div > div > textarea:focus {
        border-color: #4CAF50;
        box-shadow: 0 0 5px rgba(76, 175, 80, 0.5);
    }

    /* Section Header Hover Effect */
    .section-header:hover {
        color: #4CAF50;
        cursor: pointer;
        transition: color 0.3s;
    }

    /* Tooltip Styles */
    .tooltip {
        position: relative;
        display: inline-block;
    }

    .tooltip .tooltiptext {
        visibility: hidden;
        width: 220px;
        background-color: #555;
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 8px 0;
        position: absolute;
        z-index: 1;
        bottom: 125%; /* Position above */
        left: 50%;
        margin-left: -110px;
        opacity: 0;
        transition: opacity 0.3s;
    }

    .tooltip:hover .tooltiptext {
        visibility: visible;
        opacity: 1;
    }

    /* Footer Styles */
    .footer {
        position: fixed;
        left: 0;
        bottom: 0;
        width: 100%;
        background-color: #f7f9fc;
        color: #4B4B4B;
        text-align: center;
        padding: 10px 0;
        font-weight: 700;
    }

    /* Smooth Scrollbar */
    ::-webkit-scrollbar {
        width: 12px;
    }
    ::-webkit-scrollbar-track {
        background: #f1f1f1; 
    }
    ::-webkit-scrollbar-thumb {
        background: #888; 
        border-radius: 6px;
    }
    ::-webkit-scrollbar-thumb:hover {
        background: #555; 
    }

    </style>
    """, unsafe_allow_html=True)


def generate_pdf(sections):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []

    styles = getSampleStyleSheet()
    bold_style = ParagraphStyle(name='BoldStyle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12)
    question_style = ParagraphStyle(name='QuestionStyle', parent=styles['Normal'], fontSize=10)

    for i, section in enumerate(sections.values(), start=1):
        section_info = f'Section {chr(64 + i)}: {section["title"]} - Attempt {section["attempt_desc"]} each carrying {section["marks"]} marks'
        elements.append(Paragraph(section_info, bold_style))

        for q in section['questions']:
            elements.append(Paragraph(q, question_style))

        elements.append(Paragraph('', styles['Normal']))  # Adding extra space

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_docx(sections):
    doc = Document()

    for i, section in enumerate(sections.values(), start=1):
        section_info = f'Section {chr(64 + i)}: {section["title"]} - Attempt {section["attempt_desc"]} each carrying {section["marks"]} marks'
        paragraph = doc.add_paragraph(section_info)
        paragraph.runs[0].bold = True  # Make the section info bold

        for q in section['questions']:
            doc.add_paragraph(q)

    return doc

def main():
    inject_custom_css()

    st.markdown("<div class='main'>", unsafe_allow_html=True)
    st.markdown("<h1 style='text-align: center;'>📄 Question Paper Generator</h1>", unsafe_allow_html=True)

    num_sections = st.number_input("Enter the number of sections:", min_value=1, step=1, key="num_sections")

    sections = {}
    valid_input = False  # Flag to track if there's at least one question
    valid_sections = True  # Flag to track if all required fields are filled

    if num_sections > 0:
        for i in range(num_sections):
            section_key = f'section_{chr(65 + i)}'

            st.markdown(f"<h3 class='section-header'>🔖 Section {chr(65 + i)}</h3>", unsafe_allow_html=True)

            with st.expander(f"🔍 Expand Section {chr(65 + i)} Details", expanded=True):
                col1, col2 = st.columns(2)

                with col1:
                    section_title = st.text_input(f"Title for Section {chr(65 + i)}", key=f"{section_key}_title")

                with col2:
                    marks = st.number_input(f"Marks per Question for Section {chr(65 + i)}", min_value=1, step=1,
                                            key=f"{section_key}_marks")

                attempt_desc = st.text_input(f"Attempt description for Section {chr(65 + i)}",
                                             placeholder="e.g., Attempt all questions or Attempt 5 questions",
                                             key=f"{section_key}_attempt_desc")

                raw_questions = st.text_area(f"Enter questions for Section {chr(65 + i)} (one per line)",
                                             height=150,
                                             key=f"{section_key}_questions")

                # Check if title and attempt description are filled
                if not section_title or not attempt_desc:
                    valid_sections = False

                raw_questions_lines = [line.strip() for line in raw_questions.split('\n') if line.strip()]
                questions = [f'Q.{j + 1}: {q}' for j, q in enumerate(raw_questions_lines)]

                if questions:
                    valid_input = True  # Set flag to true if there's at least one question

                sections[section_key] = {
                    'title': section_title,
                    'attempt_desc': attempt_desc,
                    'marks': marks,
                    'questions': questions
                }

    if st.button("🎉 Generate Paper"):
        if valid_input and valid_sections:
            pdf_buffer = generate_pdf(sections)
            doc_buffer = io.BytesIO()
            doc = generate_docx(sections)
            doc.save(doc_buffer)
            doc_buffer.seek(0)

            st.success("✅ Your question paper has been generated!")

            col1, col2 = st.columns(2)
            with col1:
                st.download_button(label="📄 Download paper in PDF",
                                   data=pdf_buffer,
                                   file_name="question_paper.pdf",
                                   mime="application/pdf")
            with col2:
                st.download_button(label="📃 Download paper in DOCX",
                                   data=doc_buffer,
                                   file_name="question_paper.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            if not valid_input:
                st.warning("⚠️ Please add at least one question in any section before generating the paper.")
            if not valid_sections:
                st.warning("⚠️ Please make sure that the title and attempt description for all sections are filled.")

    st.markdown("""
        </div>
        <div class="footer">
            Developer Yasir Arfat ❤️
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()